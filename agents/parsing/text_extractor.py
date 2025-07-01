import time
from typing import List, Dict, Any, Optional
from pathlib import Path
import fitz  # PyMuPDF
import pytesseract
from PIL import Image
from agents.base_agent import BaseAgent
from database.models import AgentResponse, TextContent, DocumentElement, ContentType, BoundingBox
from utils.logger import logger

class TextExtractorAgent(BaseAgent):
    def __init__(self):
        super().__init__(
            name="TextExtractor",
            description="Extracts text content from documents while preserving formatting and reading order"
        )
    
    async def process(self, input_data: Any, **kwargs) -> AgentResponse:
        start_time = time.time()
        
        try:
            if isinstance(input_data, str):
                # Direct file path processing
                file_path = input_data
                layout_elements = kwargs.get('layout_elements', [])
                document_id = kwargs.get('document_id', Path(file_path).stem)
                elements = await self._extract_from_file(file_path, layout_elements, document_id)
            elif isinstance(input_data, dict) and 'file_path' in input_data:
                # Processing with layout information
                file_path = input_data['file_path']
                layout_elements = input_data.get('layout_elements', [])
                document_id = input_data.get('document_id', Path(file_path).stem)
                elements = await self._extract_from_file(file_path, layout_elements, document_id)
            else:
                raise ValueError("Invalid input data format")
            
            processing_time = time.time() - start_time
            
            response = self.create_response(
                content=elements,
                response_type="text_extraction",
                sources=[file_path],
                metadata={"total_text_elements": len(elements)},
                processing_time=processing_time
            )
            
            self.log_activity(f"Extracted {len(elements)} text elements")
            return response
            
        except Exception as e:
            self.log_activity(f"Text extraction failed: {e}", "error")
            processing_time = time.time() - start_time
            
            return self.create_response(
                content=f"Text extraction failed: {str(e)}",
                response_type="error",
                processing_time=processing_time
            )
    
    async def _extract_from_file(self, file_path: str, layout_elements: List[Dict], document_id: str = None) -> List[DocumentElement]:
        self.log_activity(f"Starting text extraction from: {file_path}")
        
        if not Path(file_path).exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_extension = Path(file_path).suffix.lower()
        
        if not document_id:
            document_id = Path(file_path).stem
            
        if file_extension == '.pdf':
            return await self._extract_from_pdf(file_path, layout_elements, document_id)
        elif file_extension in ['.docx', '.doc']:
            return await self._extract_from_docx(file_path, layout_elements, document_id)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
    
    async def _extract_from_pdf(self, file_path: str, layout_elements: List[Dict], document_id: str) -> List[DocumentElement]:
        elements = []
        pdf_document = None
        
        try:
            pdf_document = fitz.open(file_path)
            
            # Group layout elements by page and filter for text elements
            text_layout_elements = {}
            for elem in layout_elements:
                if elem.get('content_type') in [ContentType.TEXT.value, ContentType.HEADER.value, ContentType.LIST.value]:
                    page_num = elem.get('bounding_box', {}).get('page_number', 0)
                    if page_num not in text_layout_elements:
                        text_layout_elements[page_num] = []
                    text_layout_elements[page_num].append(elem)
            
            for page_num in range(len(pdf_document)):
                page = pdf_document.load_page(page_num)
                page_elements = await self._extract_page_text(
                    page, page_num, document_id, 
                    text_layout_elements.get(page_num, [])
                )
                elements.extend(page_elements)
            
        except Exception as e:
            self.log_activity(f"Error extracting text from PDF: {e}", "error")
            raise
        finally:
            # Ensure PDF document is always closed
            if pdf_document is not None:
                try:
                    pdf_document.close()
                except:
                    pass
        
        return elements
    
    async def _extract_page_text(self, page, page_num: int, document_id: str, layout_elements: List[Dict]) -> List[DocumentElement]:
        elements = []
        
        try:
            if layout_elements:
                # Extract text based on layout elements
                for layout_elem in layout_elements:
                    bbox_info = layout_elem.get('bounding_box', {})
                    if not bbox_info:
                        continue
                    
                    # Extract text from bounding box
                    rect = fitz.Rect(bbox_info['x1'], bbox_info['y1'], bbox_info['x2'], bbox_info['y2'])
                    text_blocks = page.get_text("dict", clip=rect)
                    
                    text_content = await self._process_text_blocks(text_blocks)
                    if text_content.content.strip():
                        element = DocumentElement(
                            document_id=document_id,
                            content_type=ContentType(layout_elem['content_type']),
                            bounding_box=BoundingBox(
                                x1=bbox_info['x1'],
                                y1=bbox_info['y1'],
                                x2=bbox_info['x2'],
                                y2=bbox_info['y2'],
                                page_number=page_num
                            ),
                            text_content=text_content,
                            metadata={
                                "extraction_method": "layout_guided",
                                "confidence": layout_elem.get('confidence', 0.8)
                            }
                        )
                        elements.append(element)
            else:
                # Extract all text from page
                text_blocks = page.get_text("dict")
                text_content = await self._process_text_blocks(text_blocks)
                
                if text_content.content.strip():
                    element = DocumentElement(
                        document_id=document_id,
                        content_type=ContentType.TEXT,
                        text_content=text_content,
                        metadata={
                            "extraction_method": "full_page",
                            "page_number": page_num
                        }
                    )
                    elements.append(element)
            
        except Exception as e:
            self.log_activity(f"Error extracting text from page {page_num}: {e}", "error")
            raise
        
        return elements
    
    async def _process_text_blocks(self, text_blocks: Dict) -> TextContent:
        text_parts = []
        font_sizes = []
        font_families = []
        is_bold = False
        is_italic = False
        
        try:
            for block in text_blocks.get("blocks", []):
                if "lines" not in block:
                    continue
                
                for line in block["lines"]:
                    line_text = ""
                    for span in line.get("spans", []):
                        text = span.get("text", "")
                        if text.strip():
                            line_text += text
                            
                            # Collect font information
                            font_sizes.append(span.get("size", 12))
                            font_families.append(span.get("font", ""))
                            
                            # Check for bold/italic
                            flags = span.get("flags", 0)
                            if flags & 2**4:  # Bold
                                is_bold = True
                            if flags & 2**1:  # Italic
                                is_italic = True
                    
                    if line_text.strip():
                        text_parts.append(line_text)
            
            # Combine text parts
            full_text = "\n".join(text_parts)
            
            # Calculate average font size
            avg_font_size = sum(font_sizes) / len(font_sizes) if font_sizes else 12
            
            # Get most common font family
            font_family = max(set(font_families), key=font_families.count) if font_families else None
            
            return TextContent(
                content=full_text,
                font_size=avg_font_size,
                font_family=font_family,
                is_bold=is_bold,
                is_italic=is_italic
            )
            
        except Exception as e:
            self.log_activity(f"Error processing text blocks: {e}", "error")
            return TextContent(content="")
    
    async def _extract_from_docx(self, file_path: str, layout_elements: List[Dict], document_id: str) -> List[DocumentElement]:
        elements = []
        
        try:
            from docx import Document
            
            doc = Document(file_path)
            
            # Extract paragraphs
            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    # Analyze paragraph formatting
                    runs = paragraph.runs
                    is_bold = any(run.bold for run in runs if run.bold is not None)
                    is_italic = any(run.italic for run in runs if run.italic is not None)
                    
                    # Determine content type based on style
                    content_type = ContentType.TEXT
                    if paragraph.style.name.startswith('Heading'):
                        content_type = ContentType.HEADER
                    elif paragraph.style.name in ['List Paragraph', 'List']:
                        content_type = ContentType.LIST
                    
                    text_content = TextContent(
                        content=paragraph.text,
                        is_bold=is_bold,
                        is_italic=is_italic,
                        reading_order=i
                    )
                    
                    element = DocumentElement(
                        document_id=document_id,
                        content_type=content_type,
                        text_content=text_content,
                        metadata={
                            "extraction_method": "docx",
                            "paragraph_index": i,
                            "style": paragraph.style.name
                        }
                    )
                    elements.append(element)
            
            # Extract tables
            for i, table in enumerate(doc.tables):
                table_text = []
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    table_text.append(" | ".join(row_text))
                
                if table_text:
                    text_content = TextContent(
                        content="\n".join(table_text),
                        reading_order=len(elements)
                    )
                    
                    element = DocumentElement(
                        document_id=document_id,
                        content_type=ContentType.TABLE,
                        text_content=text_content,
                        metadata={
                            "extraction_method": "docx_table",
                            "table_index": i,
                            "rows": len(table.rows),
                            "columns": len(table.columns) if table.rows else 0
                        }
                    )
                    elements.append(element)
            
        except Exception as e:
            self.log_activity(f"Error extracting text from DOCX: {e}", "error")
            raise
        
        return elements